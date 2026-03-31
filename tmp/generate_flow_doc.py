import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_flow_document():
    doc = Document()

    # Title
    title = doc.add_heading('Decentralized Voting System Flow', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('A comprehensive guide on the end-to-end operation of the VoteChain Decentralized Application.')

    # Introduction
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph(
        "The Decentralized Voting System is a hybrid application blending off-chain metadata storage (MongoDB) "
        "with an on-chain immutable ledger (Ethereum Smart Contract). The system ensures transparency, "
        "tamper-proof record-keeping, and cryptographic security."
    )

    # Step 1: Authentication
    doc.add_heading('2. Step 1: User Authentication & Authorization', level=1)
    doc.add_paragraph(
        "Flow:\n"
        "- Users navigate to the Decentralized Voting Web Application.\n"
        "- Users click on 'Connect Wallet' and authenticate using a Web3 wallet (e.g., MetaMask) via Sign-In-With-Ethereum (SIWE).\n"
        "- The frontend requests a cryptographic 'nonce' from the Flask Backend.\n"
        "- The user signs a specific message containing this nonce with their Ethereum private key.\n"
        "- The backend verifies the signature. If valid, a JSON Web Token (JWT) is issued with a role ('admin' or 'voter') based on their address.\n"
    )

    # Step 2: Voter Registration
    doc.add_heading('3. Step 2: Voter Registration (Admin Approval)', level=1)
    doc.add_paragraph(
        "Flow:\n"
        "- A new user requests voter registration via the frontend, submitting their details (e.g., ID, name) to the backend database.\n"
        "- The Admin views pending 'Voter Requests' in the Admin Dashboard.\n"
        "- The Admin reviews the submitted details. Upon approval, the Admin's wallet triggers the 'registerVoter' function on the Smart Contract.\n"
        "- The Smart Contract marks the user's Ethereum address as 'isRegistered' and emits a 'VoterRegistered' event.\n"
    )

    # Step 3: Election Creation Setup
    doc.add_heading('4. Step 3: Election Initialization', level=1)
    doc.add_paragraph(
        "Flow:\n"
        "- The Admin navigates to the 'Create Election' panel.\n"
        "- Admin specifies the election Title, Start Time, and End Time.\n"
        "- Admin submits a transaction to the Smart Contract's 'createElection' function.\n"
        "- Once the election is created, the Admin adds candidates by calling the 'addCandidate' function, assigning candidates to the specific Election ID.\n"
        "- The Smart Contract emits 'ElectionCreated' and 'CandidateAdded' events, synchronizing the blockchain state.\n"
    )

    # Step 4: The Voting Phase
    doc.add_heading('5. Step 4: Casting a Vote', level=1)
    doc.add_paragraph(
        "Flow:\n"
        "- An election enters the 'Active' phase (current time is between Start Time and End Time).\n"
        "- Registered voters browse the 'Active Elections' on the home page and click 'Vote Now'.\n"
        "- The voter selects their preferred candidate and confirms the transaction in their wallet.\n"
        "- The transaction is sent to the Smart Contract's 'castVote' function.\n"
        "- The contract verifies:\n"
        "  a) Admin-registered status (onlyRegistered modifier)\n"
        "  b) Election active status (electionActive modifier)\n"
        "  c) No double voting (!hasVoted check)\n"
        "- The candidate's vote count is mathematically incremented on-chain, and a 'VoteCast' event is recorded permanently.\n"
    )

    # Step 5: Tallying and Conclusion
    doc.add_heading('6. Step 5: Results & Election Conclusion', level=1)
    doc.add_paragraph(
        "Flow:\n"
        "- Users and guests can view real-time, immutable results fetched from the Smart Contract's 'getResults' function.\n"
        "- Once the election 'End Time' has passed, the Admin triggers the 'endElection' function on the Smart Contract.\n"
        "- The election state is set to inactive (active = false), preventing any further votes.\n"
        "- Final results are locked transparently on the Ethereum blockchain for auditing.\n"
    )

    # Conclusion
    doc.add_heading('7. Summary of Architecture', level=1)
    doc.add_paragraph(
        "- Frontend: Next.js + React + Tailwind CSS (Interacts with Web3/Ethers.js)\n"
        "- Backend/API: Flask Python backend (MongoDB for requests, SIWE Auth, JWTs)\n"
        "- Blockchain: Solidity Smart Contract deployed on Sepolia Testnet\n"
    )

    doc_path = os.path.join(os.path.dirname(__file__), "Decentralized_Voting_System_Flow.docx")
    doc.save(doc_path)
    print(f"Document successfully generated at: {doc_path}")

if __name__ == '__main__':
    create_flow_document()
