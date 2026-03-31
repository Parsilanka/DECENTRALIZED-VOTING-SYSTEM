import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_detailed_flow_document():
    doc = Document()
    
    styles = doc.styles
    h1 = styles['Heading 1']
    h1.font.size = Pt(16)
    h1.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    # Title
    title = doc.add_heading('Deep-Dive: Decentralized Voting System Workflow', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        'A highly technical and exhaustive end-to-end architectural flow of the VoteChain Decentralized Application, outlining '
        'precise technical state changes across the frontend, backend, and smart contract layers.'
    )

    doc.add_heading('1. Technical Stack Architecture', level=1)
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('Frontend:').bold = True
    ul.add_run(' Next.js (React), TailwindCSS, Ethers.js')
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('Backend:').bold = True
    ul.add_run(' Flask (Python), PyJWT, Flask-Limiter')
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('Database:').bold = True
    ul.add_run(' MongoDB (Off-chain rapid-access state and metadata)')
    
    ul = doc.add_paragraph(style='List Bullet')
    ul.add_run('Blockchain:').bold = True
    ul.add_run(' Solidity Smart Contract (Voting.sol) on Ethereum (Sepolia Testnet)')

    doc.add_heading('2. Phase 1: Authentication & Authorization (Sign-In-With-Ethereum)', level=1)
    doc.add_paragraph('The platform uses SIWE (EIP-4361 standard) to securely authenticate users without passwords, mapping an Ethereum public address to a role-based session.')
    doc.add_heading('Sequence of Operations:', level=2)
    seq = [
        "Initial Connection: User visits index. In WalletContext.tsx, window.ethereum.request({ method: 'eth_accounts' }) checks for an injected Web3 provider (MetaMask).",
        "Challenge Generation: User clicks Connect. The frontend invokes api.get('/api/auth/challenge/<address>'). The Flask backend generates a uuid4 nonce, temporarily caching it in memory mapped to the address.",
        "Cryptographic Signing: The frontend uses ethers.BrowserProvider. signer.signMessage(message) prompts the user to sign the challenge using their wallet's private key (ECDSA).",
        "Signature Verification & JWT: Frontend POSTs payload to /api/auth/login. Backend uses web3.py (recover_message) to mathematically derive the address. If matching, a JSON Web Token (JWT) is issued containing {address, exp, role}.",
        "Session Bootstrap: JWT is persisted in frontend localStorage. WalletContext state is updated."
    ]
    for idx, s in enumerate(seq, 1):
        doc.add_paragraph(f"{idx}. {s}", style='List Number')

    doc.add_heading('3. Phase 2: Voter Registration Workflow', level=1)
    doc.add_paragraph('Public blockchain applications require "Permissioned" access. This phase prevents Sybil attacks by requiring the Admin to formally whitelist an address on-chain.')
    doc.add_heading('Sequence of Operations:', level=2)
    seq2 = [
        "Application Submission: Voter submits details to /api/voters/apply. Writes to MongoDB voter_requests stringifying status as 'pending'.",
        "Admin Review: Admin accesses /admin/voters. Backend validates @token_required decorators and 'admin' JWT role mapping.",
        "Approval Execution: Admin clicks Approve. Backend mutates Mongo document status to 'approving' synchronously, then spawns a background threading.Thread.",
        "On-Chain Mutation: The thread constructs a transaction calling registerVoter(address). It signs using the Admin Private Key. Contract modifier onlyAdmin accepts. The contract sets isRegistered[_voter] = true and emits VoterRegistered event."
    ]
    for idx, s in enumerate(seq2, 1):
        doc.add_paragraph(f"{idx}. {s}", style='List Number')

    doc.add_heading('4. Phase 3: Election and Candidate Setup', level=1)
    doc.add_heading('Sequence of Operations:', level=2)
    seq3 = [
        "Election Initialization: Admin submits Title and Epoch Times to /api/admin/election. A thread hits contract createElection(). The contract evaluates time constraints and pushes to an Active Elections mapping.",
        "Candidate Initialization: Admin submits Candidate data targeting an election_id. Contract addCandidate() asserts elections[_electionId].active == true. Backend stores heavy strings (Manifesto, Image URLs) in MongoDB indexing on election_id and candidate_id to save GAS fees."
    ]
    for idx, s in enumerate(seq3, 1):
        doc.add_paragraph(f"{idx}. {s}", style='List Number')

    doc.add_heading('5. Phase 4: Constructing & Casting the Vote', level=1)
    doc.add_paragraph('To guarantee strict decentralization, the frontend interacts directly with the active Smart Contract bypassing backend intermediaries for execution.')
    doc.add_heading('Sequence of Operations:', level=2)
    seq4 = [
        "Pre-Vote Checks: Frontend queries /api/voters/has-voted/{id}/{address} visually disabling the button if user previously voted.",
        "Payload Construction: Frontend builds castVoteAction via the local ethers.Contract object, dynamically linking the user signer.",
        "Wallet Interaction & Gas: MetaMask intercepts the contract.castVote() transaction. The user spends their own ETH (Gas) signing the transaction directly to Sepolia mempool.",
        "Contract Validation (On EVM): The Block processes. The contract validates nonReentrant, onlyRegistered checks if Admin verified them earlier, electionActive evaluates current timestamp, and !hasVoted checks to avoid double-voting.",
        "State Finalization: voteCount increments mathematically on-chain. VoteCast event fires. Frontend await tx.wait() receives receipt and returns the Transaction Hash."
    ]
    for idx, s in enumerate(seq4, 1):
        doc.add_paragraph(f"{idx}. {s}", style='List Number')

    doc.add_heading('6. Phase 5: Result Tallying & Conclusion', level=1)
    doc.add_heading('Sequence of Operations:', level=2)
    seq5 = [
        "Direct On-chain Reading: Backend calls getResults() mapping accurate on-chain voteCounts and joining with MongoDB imagery/bios silently for REST output.",
        "Natural End: If EVM block timestamp strictly exceeds endTime, contract natively drops any castVote attempts causing reverts.",
        "Manual Locking: Admin triggers /api/admin/end-election/[id]. The server initiates endElection() Smart Contract state modification, emitting ElectionEnded and forcefully collapsing the boolean variable active = false permanently."
    ]
    for idx, s in enumerate(seq5, 1):
        doc.add_paragraph(f"{idx}. {s}", style='List Number')

    doc_path = os.path.join(os.path.dirname(__file__), "Decentralized_Voting_System_Detailed_Flow.docx")
    doc.save(doc_path)
    print(f"Detailed Document successfully generated at: {doc_path}")

if __name__ == '__main__':
    create_detailed_flow_document()
